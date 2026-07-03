# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("output/docx/relatorio_mochila_paa.docx")


def set_run_font(run, size=12, bold=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc, text="", bold_prefix=None, align=None):
    paragraph = doc.add_paragraph()
    if align:
        paragraph.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(run, bold=True)
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=12, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def make_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "F2F2F2")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def build_document():
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)

    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.15
    for heading_style in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[heading_style]
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Análise de Algoritmos para o Problema da Mochila 0-1 com Peso e Volume")
    set_run_font(run, bold=True)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run("Diogo Tadeu, Thiago Henrique, Wendel Cauã, Marcelo")
    set_run_font(run)

    add_heading(doc, "Resumo", 1)
    add_paragraph(doc, "Este relatório apresenta a implementação e avaliação de três estratégias para o problema da Mochila 0-1 com duas restrições: peso e volume. Foram implementados Backtracking com podas, Programação Dinâmica com estados esparsos e Branch and Bound com limites superiores. Os experimentos avaliaram o tempo de execução em 20.113 execuções, distribuídas em 867 combinações de quantidade de itens, capacidade de peso e capacidade de volume. Os resultados mostram que a Programação Dinâmica foi a alternativa mais estável nas instâncias analisadas, enquanto o Backtracking apresentou os maiores picos de tempo. O Branch and Bound reduziu bastante o custo em relação ao Backtracking, mas ainda apresentou crescimento elevado em instâncias difíceis.")

    add_heading(doc, "1. Introdução", 1)
    add_paragraph(doc, "O problema da Mochila 0-1 consiste em selecionar um subconjunto de itens de forma a maximizar o lucro total, respeitando restrições de capacidade. Na versão estudada neste trabalho, cada item possui três atributos inteiros positivos: lucro, peso e volume. A mochila possui uma capacidade máxima de peso W e uma capacidade máxima de volume V. Como a escolha é 0-1, cada item pode ser escolhido no máximo uma vez.")
    add_paragraph(doc, "O objetivo do trabalho foi implementar e comparar três algoritmos exatos para resolver o problema: Backtracking, Programação Dinâmica e Branch and Bound. A comparação foi feita principalmente pela métrica de tempo de execução, considerando diferentes tamanhos de entrada e diferentes capacidades de peso e volume. Além disso, foram produzidos arquivos CSV, análises estatísticas, gráficos e dashboards para apoiar a interpretação dos resultados.")
    add_paragraph(doc, "De forma resumida, os resultados indicaram que a Programação Dinâmica apresentou maior estabilidade, principalmente por reutilizar estados já calculados. O Backtracking foi fortemente afetado pelo aumento do número de itens. O Branch and Bound teve desempenho intermediário: melhor que o Backtracking em muitos casos, mas ainda sensível à qualidade das podas. O restante do relatório está organizado em descrição dos algoritmos, análise de complexidade, avaliação experimental, conclusão e referências.")

    add_heading(doc, "2. Descrição dos algoritmos e complexidades", 1)
    add_heading(doc, "2.1 Backtracking com poda", 2)
    add_paragraph(doc, "O Backtracking percorre uma árvore de decisão em que, para cada item, existem duas possibilidades: incluir ou não incluir o item na mochila. A implementação ordena os itens por uma densidade normalizada, considerando peso e volume relativos, e usa duas podas principais: descarte de ramos que excedem peso ou volume e descarte de ramos cujo lucro atual somado ao lucro restante possível não supera a melhor solução encontrada.")
    add_paragraph(doc, "No pior caso, as podas podem não eliminar ramos suficientes, e o algoritmo ainda precisa analisar um número exponencial de subconjuntos. Assim, a complexidade de tempo no pior caso é O(2^n), em que n é a quantidade de itens. A complexidade de espaço é O(n), considerando a pilha recursiva e o vetor de itens escolhidos no caminho atual.")

    add_heading(doc, "2.2 Programação Dinâmica com estados esparsos", 2)
    add_paragraph(doc, "A Programação Dinâmica utilizada é iterativa, não recursiva com memoization. O algoritmo mantém estados identificados pelo par (peso, volume), armazenando o maior lucro encontrado para cada combinação alcançável. Em vez de construir uma matriz tridimensional completa dp[i][w][v], a implementação usa uma tabela hash, por meio de unordered_map, para guardar apenas estados alcançados. Também são removidos estados dominados, isto é, estados que possuem peso e volume não melhores e lucro menor ou igual a outro estado.")
    add_paragraph(doc, "A formulação clássica da mochila 0-1 com duas restrições possui tempo O(nWV) e espaço O(WV), usando otimização por camadas, ou O(nWV) se toda a matriz tridimensional for armazenada. Na implementação esparsa, o custo prático depende da quantidade S de estados mantidos. Assim, o tempo observado pode ser descrito como O(nS), mais o custo das podas de dominância, e o espaço como O(S). No pior caso, S pode se aproximar de W·V, mantendo a natureza pseudo-polinomial do método.")

    add_heading(doc, "2.3 Branch and Bound", 2)
    add_paragraph(doc, "O Branch and Bound também explora uma árvore de decisão 0-1, mas usa limites superiores para evitar a exploração de ramos que não podem melhorar a melhor solução conhecida. A implementação começa com uma solução inicial gulosa, o que aumenta o poder das podas desde o início. O limite superior combina três estimativas: valor restante máximo, limite fracionário por peso e limite fracionário por volume. A solução final continua sendo 0-1; a ideia fracionária é usada apenas para estimar um limite superior.")
    add_paragraph(doc, "No pior caso, o Branch and Bound ainda pode visitar uma quantidade exponencial de nós. Como o cálculo do limite pode custar tempo proporcional ao número de itens restantes, a complexidade de tempo no pior caso é O(n·2^n). Na prática, as podas costumam reduzir bastante o número de nós visitados. A complexidade de espaço é O(n), considerando a profundidade da recursão e os vetores de itens escolhidos, além das listas auxiliares ordenadas.")

    add_heading(doc, "3. Avaliação experimental", 1)
    add_heading(doc, "3.1 Configuração dos experimentos", 2)
    add_paragraph(doc, "Os programas foram implementados em C++ e compilados por Make para execução no WSL. Foram gerados três executáveis separados, um para cada algoritmo, além de um gerador automático de instâncias e um executor de experimentos. As instâncias possuem itens com valores inteiros positivos, e os índices dos itens começam em 1. Para cada combinação solicitada de n, W e V, foram geradas repetições aleatórias, permitindo calcular médias e desvios.")
    add_paragraph(doc, "A saída dos experimentos foi registrada em CSV com, entre outros campos, quantidade de itens, capacidades, repetição, algoritmo, lucro obtido, itens escolhidos, tempo em segundos e status. Foi adotado limite de execução de 7200 segundos por algoritmo, permitindo interromper casos que ultrapassassem duas horas e prosseguir com o restante dos testes.")

    add_heading(doc, "3.2 Métrica de avaliação", 2)
    add_paragraph(doc, "A principal métrica de avaliação foi o tempo de execução em segundos. Também foram registrados lucro máximo e itens escolhidos para verificar a consistência das soluções. Como os algoritmos são exatos, espera-se que, quando finalizados corretamente, retornem o mesmo lucro ótimo para a mesma instância. Para comparação estatística, foram aplicados testes por combinação de n, W e V. Quando havia três algoritmos, foi usado o teste de Friedman; nas comparações par a par, foi usado Wilcoxon. Considerou-se diferença estatisticamente significativa quando p < 0,05 e empate estatístico quando p >= 0,05.")

    add_heading(doc, "3.3 Resultados obtidos", 2)
    add_paragraph(doc, "Ao todo, foram analisadas 20.113 execuções em 867 combinações de parâmetros. A Tabela 1 resume os tempos médios e máximos por conjunto de experimentos e algoritmo.")
    make_table(
        doc,
        ["CSV", "Alg.", "Exec.", "Média (s)", "Máx. (s)"],
        [
            ["100a600", "DP", "1403", "0,25", "5,27"],
            ["100a600", "BT", "1401", "59,94", "2686,18"],
            ["100a600", "B&B", "1401", "5,59", "306,02"],
            ["300 itens + W/V", "DP", "1080", "1,14", "76,72"],
            ["300 itens + W/V", "BT", "1080", "0,80", "31,85"],
            ["300 itens + W/V", "B&B", "1080", "0,11", "3,66"],
            ["B&B vs DP 1000", "DP", "5900", "0,59", "9,85"],
            ["B&B vs DP 1000", "B&B", "5900", "11,40", "1362,08"],
            ["variação total", "DP", "291", "0,018", "0,20"],
            ["variação total", "BT", "288", "260,06", "13565,98"],
            ["variação total", "B&B", "289", "17,81", "631,14"],
        ],
        widths=[2.0, 0.7, 0.75, 0.9, 0.95],
    )
    add_paragraph(doc, "Os resultados mostram grande diferença entre os algoritmos. No conjunto 100a600, por exemplo, o Backtracking teve média de 59,94 s, enquanto o Branch and Bound teve 5,59 s e a Programação Dinâmica 0,25 s. No conjunto variação total, o Backtracking atingiu média de 260,06 s e máximo superior a 13.565 s, evidenciando o crescimento exponencial em instâncias difíceis.")
    add_paragraph(doc, "A Programação Dinâmica apresentou comportamento mais previsível e tempos menores em grande parte dos testes. Isso ocorreu porque o método reutiliza estados e evita enumerar subconjuntos completos. O Branch and Bound reduziu substancialmente o custo em relação ao Backtracking, mas ainda teve picos altos quando os limites superiores não foram suficientemente fortes para eliminar muitos ramos. Em alguns conjuntos com capacidades muito grandes, a Programação Dinâmica sentiu mais o aumento da quantidade de estados possíveis, o que explica casos em que B&B foi competitivo ou mais rápido.")
    add_paragraph(doc, "Nos testes estatísticos, a maioria das comparações indicou diferença significativa de tempo entre os algoritmos. As comparações DP vs BT apresentaram 249 diferenças significativas e 26 empates estatísticos; DP vs B&B apresentou 733 diferenças e 132 empates; BT vs B&B apresentou 274 diferenças e apenas 1 empate. O teste de Friedman, aplicado quando havia os três algoritmos, indicou diferença significativa em 275 combinações analisadas.")

    add_heading(doc, "4. Discussão", 1)
    add_paragraph(doc, "O comportamento observado está de acordo com as complexidades teóricas. O Backtracking possui pior caso exponencial, e seus tempos cresceram muito quando n aumentou. As podas simples ajudam, mas não eliminam a explosão combinatória. O Branch and Bound também possui pior caso exponencial, porém os limites superiores permitem cortar subárvores inteiras, explicando sua melhora em relação ao Backtracking. Ainda assim, quando muitos ramos parecem promissores, o algoritmo continua explorando uma quantidade elevada de nós.")
    add_paragraph(doc, "A Programação Dinâmica tem comportamento pseudo-polinomial e depende das capacidades W e V e da quantidade de estados mantidos. A versão esparsa implementada reduziu o consumo de memória e o tempo em relação a uma matriz tridimensional completa, especialmente quando poucas combinações de peso e volume permaneceram relevantes. Por isso, nos testes com instâncias grandes, a DP foi geralmente a opção mais estável.")

    add_heading(doc, "5. Conclusão", 1)
    add_paragraph(doc, "O trabalho implementou e comparou três algoritmos exatos para a Mochila 0-1 com restrições de peso e volume. A avaliação experimental mostrou que a escolha do algoritmo depende fortemente das características da instância. Para os testes realizados, a Programação Dinâmica com estados esparsos foi a estratégia mais estável e apresentou os menores tempos em grande parte dos casos. O Branch and Bound foi uma alternativa importante por reduzir bastante o custo do Backtracking, mas ainda apresentou picos de tempo em instâncias difíceis. O Backtracking, apesar de simples e útil para compreender a estrutura combinatória do problema, foi o menos adequado para instâncias grandes.")
    add_paragraph(doc, "Como trabalhos futuros, seria possível medir também o número de nós visitados e o número de estados mantidos, além de comparar diferentes estratégias de ordenação, limites superiores e podas de dominância. Também seria interessante avaliar o consumo de memória de forma sistemática, já que esse fator é especialmente relevante para a Programação Dinâmica.")

    add_heading(doc, "Referências bibliográficas", 1)
    for reference in [
        "CORMEN, T. H.; LEISERSON, C. E.; RIVEST, R. L.; STEIN, C. Introduction to Algorithms. 3. ed. Cambridge: MIT Press, 2009.",
        "MARTELLO, S.; TOTH, P. Knapsack Problems: Algorithms and Computer Implementations. Chichester: John Wiley & Sons, 1990.",
        "KELLERER, H.; PFERSCHY, U.; PISINGER, D. Knapsack Problems. Berlin: Springer, 2004.",
        "SIPSER, M. Introduction to the Theory of Computation. 3. ed. Boston: Cengage Learning, 2012.",
    ]:
        add_paragraph(doc, reference)

    props = doc.core_properties
    props.title = "Análise de Algoritmos para o Problema da Mochila 0-1 com Peso e Volume"
    props.author = "Diogo Tadeu; Thiago Henrique; Wendel Cauã; Marcelo"
    props.subject = "Trabalho Prático de PAA"
    props.keywords = "mochila 0-1; programação dinâmica; backtracking; branch and bound; PAA"

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
